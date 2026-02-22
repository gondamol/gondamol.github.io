import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Color blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def md_to_docx(md_file_path, docx_file_path):
    document = Document()
    
    # Set margins to be a bit narrower for a "resume" feel
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        if not line:
            if in_table:
                 # End of table
                 pass
            else:
                 document.add_paragraph()
            continue
            
        # Headings
        if line.startswith('# '):
            h = document.add_heading(level=0)
            run = h.add_run(line[2:])
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        elif line.startswith('## '):
            h = document.add_heading(level=1)
            line_content = line[3:].replace('**', '')
            run = h.add_run(line_content)
            run.bold = True
            # Unique GD Blue color
            run.font.color.rgb = RGBColor(0, 51, 102) 
            continue
        elif line.startswith('### '):
            h = document.add_heading(level=2)
            line_content = line[4:].replace('**', '')
            run = h.add_run(line_content)
            run.bold = True
            run.font.color.rgb = RGBColor(50, 50, 50)
            continue

        # Horizontal Rule
        if line.startswith('---'):
             p = document.add_paragraph()
             run = p.add_run('________________________________________________________________________________')
             run.font.color.rgb = RGBColor(200, 200, 200)
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             continue

        # Table
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            row_content = [cell.strip() for cell in line.strip('|').split('|')]
            if '---' in row_content[0]:
                continue
            table_data.append(row_content)
            
            # Flush table immediately if next line is empty or not a table row
            # Actually, let's just keep appending and flush when non-table line hits
            continue
        else:
            if in_table:
                if table_data:
                    rows = len(table_data)
                    cols = max(len(r) for r in table_data)
                    table = document.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    for i, row in enumerate(table_data):
                        cells = table.rows[i].cells
                        for j, cell_text in enumerate(row):
                             if j < len(cells):
                                # Clean bold markers
                                clean_text = cell_text.replace('**', '')
                                cells[j].text = clean_text
                                # Bold headers
                                if i == 0:
                                    for paragraph in cells[j].paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                                            run.font.color.rgb = RGBColor(0, 0, 0)
                                            # Light gray shading for header? python-docx complex shading omitted for safety
                in_table = False
                table_data = []

        # List items
        if line.startswith('* ') or line.startswith('- '):
            p = document.add_paragraph(style='List Bullet')
            clean_line = line[2:]
        else:
            p = document.add_paragraph()
            clean_line = line

        # Process formatted text (Bold and Links)
        # Regex to split by links first: [text](url)
        parts = re.split(r'(\[.*?\]\(.*?\))', clean_line)
        for part in parts:
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if link_match:
                text, url = link_match.groups()
                # Determine if link text was intended to be bold
                if text.startswith('**') and text.endswith('**'):
                    text = text[2:-2]
                try:
                    add_hyperlink(p, url, text)
                except:
                    # Fallback if XML manipulation fails
                    run = p.add_run(text)
                    run.font.color.rgb = RGBColor(0, 0, 255)
            else:
                # Process Bold: **text**
                subparts = re.split(r'(\*\*.*?\*\*)', part)
                for subpart in subparts:
                    if subpart.startswith('**') and subpart.endswith('**'):
                         p.add_run(subpart[2:-2]).bold = True
                    else:
                         p.add_run(subpart)
    
    # Final flush if table was last
    if in_table and table_data:
        rows = len(table_data)
        cols = max(len(r) for r in table_data)
        table = document.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for i, row in enumerate(table_data):
            cells = table.rows[i].cells
            for j, cell_text in enumerate(row):
                 if j < len(cells):
                    cells[j].text = cell_text.replace('**', '')

    document.save(docx_file_path)

if __name__ == "__main__":
    # Resume
    md_resume = r"c:\Users\HFD 2\pc\JBS\2026\feb\GiveDirectly\Nichodemus_Amollo_GD_Resume.md"
    docx_resume = r"c:\Users\HFD 2\pc\JBS\2026\feb\GiveDirectly\Nichodemus_Amollo_GD_Resume.docx"
    try:
        md_to_docx(md_resume, docx_resume)
        print("Resume converted.")
    except Exception as e:
        print(f"Resume Error: {e}")

    # Cover Letter
    md_cl = r"c:\Users\HFD 2\pc\JBS\2026\feb\GiveDirectly\Nichodemus_Amollo_GD_Cover_Letter.md"
    docx_cl = r"c:\Users\HFD 2\pc\JBS\2026\feb\GiveDirectly\Nichodemus_Amollo_GD_Cover_Letter.docx"
    try:
        md_to_docx(md_cl, docx_cl)
        print("Cover Letter converted.")
    except Exception as e:
        print(f"Cover Letter Error: {e}")
