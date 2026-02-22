import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_hyperlink(paragraph, url, text):
    """
    Simulated hyperlink: visually formatted text.
    Actually adding a real hyperlink is complex and version-dependent in python-docx.
    For now, we just make it blue and underlined.
    """
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.underline = True
    return run

def md_link_to_text(match):
    text = match.group(1)
    url = match.group(2)
    return text, url

def md_to_docx(md_file_path, docx_file_path):
    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Headings
        if line.startswith('# '):
            h = document.add_heading(level=0)
            run = h.add_run(line[2:])
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        elif line.startswith('## '):
            h = document.add_heading(level=1)
            line_content = line[3:].replace('**', '')
            run = h.add_run(line_content)
            run.bold = True
            run.font.color.rgb = RGBColor(46, 116, 181)
            continue
        elif line.startswith('### '):
            h = document.add_heading(level=2)
            line_content = line[4:].replace('**', '')
            run = h.add_run(line_content)
            run.bold = True
            continue

        # Horizontal Rule
        if line.startswith('---'):
             document.add_paragraph()
             continue

        # Table
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            row_content = [cell.strip() for cell in line.strip('|').split('|')]
            if '---' in row_content[0]:
                continue
            table_data.append(row_content)
            continue
        else:
            if in_table:
                if table_data:
                    rows = len(table_data)
                    cols = max(len(r) for r in table_data)
                    table = document.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    for i, row in enumerate(table_data):
                        for j, cell_text in enumerate(row):
                            if j < len(table.rows[i].cells):
                                cell = table.rows[i].cells[j]
                                cell.text = cell_text.replace('**', '')
                in_table = False
                table_data = []

        # List items
        if line.startswith('* ') or line.startswith('- '):
            p = document.add_paragraph(style='List Bullet')
            clean_line = line[2:]
        else:
            p = document.add_paragraph()
            clean_line = line

        # Parse Bold and Links
        # We need to process the line character by character or use complex regex to handle both.
        # Simple approach: Handle links first, then bold?
        # Links are [text](url). Bold is **text**.
        # If we just replace [text](url) with text, we lose the URL.
        # Let's iterate with a tokenizer of sorts.
        
        # Split by links: [text](url)
        # Regex for link: \[([^\]]+)\]\(([^)]+)\)
        
        parts = re.split(r'(\[.*?\]\(.*?\))', clean_line)
        for part in parts:
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if link_match:
                # It's a link
                text, url = link_match.groups()
                # Check if text is bolded
                if text.startswith('**') and text.endswith('**'):
                    text = text[2:-2]
                    # Add bold link?
                    run = p.add_run(text)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True
                else:
                    add_hyperlink(p, url, text)
            else:
                # It's text, identifying bold **text**
                subparts = re.split(r'(\*\*.*?\*\*)', part)
                for subpart in subparts:
                    if subpart.startswith('**') and subpart.endswith('**'):
                        p.add_run(subpart[2:-2]).bold = True
                    else:
                        p.add_run(subpart)

    document.save(docx_file_path)

if __name__ == "__main__":
    md_file = r"c:\Users\HFD 2\pc\JBS\2026\feb\Educate\Nichodemus_Amollo_Educate_Final_Resume.md"
    docx_file = r"c:\Users\HFD 2\pc\JBS\2026\feb\Educate\Nichodemus_Amollo_Educate_Final_Resume.docx"
    try:
        md_to_docx(md_file, docx_file)
        print("Successfully converted.")
    except Exception as e:
        print(f"Error: {e}")
