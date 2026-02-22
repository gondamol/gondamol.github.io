from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_resume():
    document = Document()

    # Adjust margins - Standard professional
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # --- CONTENT ---
    
    # Header
    name = "NICHODEMUS AMOLLO"
    title_line = "Staff Data Analyst | Strategic Data Partner"
    contact_line = "Nairobi, Kenya | +254 725 737 867 | nichodemuswerre@gmail.com"
    portfolio_line = "GitHub: github.com/gondamol | LinkedIn: linkedin.com/in/nichodemusamollo"

    summary = (
        "Senior Data Analyst with over 8 years of experience transforming complex, unstructured datasets into "
        "strategic business intelligence. Proven track record of influencing executive decision-making through rigorous "
        "data storytelling and predictive modeling. Expert in advanced SQL, Python, and BI dashboard design (Tableau, "
        "Power BI, Looker Studio), having led large-scale data system migrations and fostered data literacy across "
        "multi-country organizations. Adept at navigating highly regulated environments (health/development) and applying "
        "those frameworks to drive scalable, compliant growth. Passionate about mentoring junior analysts and building "
        "robust, reusable data products that serve as force multipliers for the business."
    )

    core_competencies = [
        ("Data Strategy & Influence", "Executive Dashboarding, KPI Definition, Board-Level Reporting, Strategic Planning"),
        ("Advanced Analytics", "Expert SQL (Window Functions, CTEs), Data Modeling, ETL Pipelines, Predictive Analysis"),
        ("Visualization & BI", "Tableau, Power BI, Looker Studio, Streamlit (Python), Superset"),
        ("Technical Leadership", "Mentoring Junior Analysts, Code Reviews, Architectural Guidance, Data Governance"),
        ("Cloud & Infrastructure", "AWS Certified Cloud Practitioner, dbt Concepts, Snowflake Familiarity, Structured Databases")
    ]

    experience = [
        {
            "role": "Lead Research Data Analyst & Strategist",
            "company": "Georgetown University East Africa",
            "location_date": "Remote / Regional (Kenya based) | Apr 2025 – Present",
            "details": [
                "**Strategic Data Leadership:** Advise senior country directors and international stakeholders on program performance, leveraging data insights to optimize multi-million dollar portfolio allocation across East Africa.",
                "**BI System Architecture:** Architecting the migration of legacy reporting systems to modern, automated dashboards (Power BI/Looker), reducing manual reporting time by 60% and improving data refresh rates.",
                "**Advanced SQL & Modeling:** Developed complex SQL queries to model health outcomes and program efficacy, directly influencing funding renewal decisions.",
                "**Mentorship & Capacity Building:** Lead a team of data associates, conducting regular code reviews, pair programming sessions, and establishing best practices for reproducible research and data quality.",
                "**Cross-Functional Collaboration:** Partner with non-technical program managers to translate ambiguous business questions into actionable analytical plans."
            ]
        },
        {
            "role": "Senior Statistician & Data Analyst",
            "company": "Kenya Medical Research Institute (KEMRI)",
            "location_date": "Kenya | Apr 2023 – Mar 2025",
            "details": [
                "**Predictive Analytics:** Designed and implemented predictive models using Python and R to forecast disease trends, enabling proactive resource allocation by national health ministries.",
                "**Executive Reporting:** Created and maintained high-visibility dashboards tracking key performance indicators (KPIs) for government-funded initiatives, ensuring alignment with national strategic goals.",
                "**Data Quality & Governance:** Established rigorous data governance frameworks for managing sensitive health data, ensuring compliance with ethical standards and regulatory requirements (paralleling compliance needs in regulated markets).",
                "**Stakeholder Management:** Regularly presented complex statistical findings to diverse audiences, including government officials and international donors, driving evidence-based policy changes."
            ]
        },
        {
            "role": "Regional Data Systems Specialist",
            "company": "VLIR-UOS / JOOUST Regional Programs",
            "location_date": "Kenya, Tanzania, Rwanda | Nov 2021 – Apr 2022",
            "details": [
                "**Data Pipeline Optimization:** Led the development of regional information management systems, harmonizing disparate data sources into a unified warehouse for cross-country analysis.",
                "**Digital Transformation:** Championed the transition from static reports to dynamic, interactive visualizations, significantly enhancing the visibility of program impact.",
                "**ETL Process Design:** Designed and maintained ETL workflows to ensure data integrity and timeliness for regional reporting."
            ]
        },
        {
            "role": "Data Analyst & MEAL Specialist",
            "company": "LERIS Hub",
            "location_date": "Kenya | Sep 2017 – May 2021",
            "details": [
                "**Impact Analysis:** Conducted rigorous quantitative analysis to evaluate program effectiveness, providing actionable recommendations that improved operational efficiency by 25%.",
                "**Data Storytelling:** Translated complex datasets into compelling narratives for donor reports and stakeholder presentations.",
                "**Tool Development:** Developed automated data collection tools and analysis scripts, streamlining the monitoring and evaluation process."
            ]
        }
    ]

    education = [
        {
            "degree": "MSc – Statistics & Data Science (Expected 2026)",
            "school": "Jaramogi Oginga Odinga University of Science & Technology",
            "detail": None 
        },
        {
            "degree": "BSc (Honours) – Statistics",
            "school": "University of Nairobi",
            "detail": None
        }
    ]

    certs = [
        "**AWS Certified Cloud Practitioner**",
        "**Google Data Analytics Professional Certificate**",
        "**Monitoring & Evaluation for Global Health** – University of Washington",
        "**Biomedical Research Ethics** – CITI Program (Data Privacy/Compliance Focus)"
    ]

    tech_skills = [
        "**Statistical & Analytical:** Python (Pandas, NumPy, Scikit-learn), R (Tidyverse, Shiny), Advanced SQL (Window Functions, CTEs), STATA.",
        "**Visualization:** Tableau, Power BI, Looker Studio, Streamlit.",
        "**Data Engineering:** Structured Databases, ETL Pipeline Concepts, API Integrations, Data Modeling.",
        "**Tools & Platforms:** Git/GitHub, AWS (Cloud Practitioner), dbt (Familiarity), Snowflake (Familiarity)."
    ]

    # --- STYLE HELPERS ---
    
    def add_section_header(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        
        # Add border below
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        pbdr = OxmlElement('w:pbdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000') # Black standard border
        pbdr.append(bottom)
        pPr.append(pbdr)

        runner = p.add_run(text.upper())
        runner.bold = True
        runner.font.size = Pt(11)
        runner.font.name = 'Calibri'
        runner.font.color.rgb = RGBColor(0, 0, 0) # Black

    # --- DOCUMENT BUILDING ---

    # HEADER
    p_name = document.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(0)
    r_name = p_name.add_run(name)
    r_name.bold = True
    r_name.font.size = Pt(22)
    r_name.font.name = 'Calibri'
    r_name.font.color.rgb = RGBColor(0, 0, 0) # Black name

    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run(title_line)
    r_title.bold = True
    r_title.font.size = Pt(11)
    r_title.font.name = 'Calibri'

    p_contact = document.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(0)
    r_contact = p_contact.add_run(contact_line)
    r_contact.font.size = Pt(10)
    r_contact.font.name = 'Calibri'

    p_port = document.add_paragraph()
    p_port.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_port.paragraph_format.space_after = Pt(12)
    r_port = p_port.add_run(portfolio_line)
    r_port.font.size = Pt(10)
    r_port.font.name = 'Calibri'
    r_port.font.color.rgb = RGBColor(0, 0, 255) # Blue link

    # PROFESSIONAL SUMMARY
    add_section_header(document, "Professional Summary")
    p_sum = document.add_paragraph(summary)
    p_sum.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_sum.paragraph_format.space_before = Pt(4)
    for run in p_sum.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # TECHNICAL SKILLS (Moved up for Analyst Role)
    add_section_header(document, "Technical Skills")
    for skill in tech_skills:
        p_skill = document.add_paragraph(style='List Bullet')
        p_skill.paragraph_format.space_after = Pt(2)
        if "**" in skill:
            parts = skill.split("**")
            for i, part in enumerate(parts):
                r = p_skill.add_run(part)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)
                if i % 2 == 1:
                    r.bold = True
        else:
            r = p_skill.add_run(skill)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

    # PROFESSIONAL EXPERIENCE
    add_section_header(document, "Professional Experience")
    
    for job in experience:
        # Job Title Line
        p_job = document.add_paragraph()
        p_job.paragraph_format.space_before = Pt(12)
        p_job.paragraph_format.space_after = Pt(0)
        
        r_role = p_job.add_run(job['role'])
        r_role.bold = True
        r_role.font.size = Pt(11)
        r_role.font.name = 'Calibri'
        
        if job['company']:
            p_job.add_run(" | ")
            r_comp = p_job.add_run(job['company'])
            r_comp.bold = True 
            r_comp.font.size = Pt(11)
            r_comp.font.name = 'Calibri'

        # Location | Date Line
        p_loc = document.add_paragraph()
        p_loc.paragraph_format.space_after = Pt(4)
        r_loc = p_loc.add_run(job['location_date'])
        r_loc.italic = True
        r_loc.font.size = Pt(10)
        r_loc.font.name = 'Calibri'

        # Bullets
        for detail in job['details']:
            p_bull = document.add_paragraph(style='List Bullet')
            p_bull.paragraph_format.space_after = Pt(2)
            
            # Formatting logic for **Bold**
            if "**" in detail:
                parts = detail.split("**")
                for i, part in enumerate(parts):
                    r = p_bull.add_run(part)
                    r.font.name = 'Calibri'
                    r.font.size = Pt(11)
                    if i % 2 == 1:
                        r.bold = True
            else:
                r = p_bull.add_run(detail)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)

    # EDUCATION
    add_section_header(document, "Education")
    for edu in education:
        p_edu = document.add_paragraph()
        p_edu.paragraph_format.space_before = Pt(6)
        p_edu.paragraph_format.space_after = Pt(0)
        
        r_deg = p_edu.add_run(edu['degree'])
        r_deg.bold = True
        r_deg.font.name = 'Calibri'
        r_deg.font.size = Pt(11)
        
        p_sch = document.add_paragraph()
        p_sch.paragraph_format.space_after = Pt(0)
        r_sch = p_sch.add_run(edu['school'])
        r_sch.font.name = 'Calibri'
        r_sch.font.size = Pt(11)
        
        if edu['detail']:
            p_det = document.add_paragraph()
            p_det.paragraph_format.space_after = Pt(2)
            r_det = p_det.add_run(edu['detail'])
            r_det.italic = True
            r_det.font.name = 'Calibri'
            r_det.font.size = Pt(10)

    # CERTIFICATIONS
    add_section_header(document, "Certifications")
    for cert in certs:
        p_cert = document.add_paragraph(style='List Bullet')
        p_cert.paragraph_format.space_after = Pt(2)
        if "**" in cert:
            parts = cert.split("**")
            for i, part in enumerate(parts):
                r = p_cert.add_run(part)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)
                if i % 2 == 1:
                    r.bold = True
        else:
            r = p_cert.add_run(cert)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)


    document.save('Nichodemus_Amollo_Weedmaps_Resume.docx')

if __name__ == "__main__":
    create_resume()
