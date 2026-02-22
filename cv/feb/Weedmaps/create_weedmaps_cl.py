from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def create_cl():
    document = Document()

    # Adjust margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Header Date
    today = datetime.date.today().strftime("%B %d, %Y")
    p_date = document.add_paragraph(today)
    p_date.paragraph_format.space_after = Pt(12)

    # Recipient
    p_recipient = document.add_paragraph()
    p_recipient.add_run("Recruitment Team\nWeedmaps\nRemote / Global")
    p_recipient.paragraph_format.space_after = Pt(12)

    # Salutation
    p_salutation = document.add_paragraph("Dear Hiring Manager,")
    p_salutation.paragraph_format.space_after = Pt(12)

    # Body Paragraphs
    p1 = document.add_paragraph()
    p1.add_run("I am writing to express my strong interest in the Staff Data Analyst position at Weedmaps. With over 8 years of experience leading complex data strategy initiatives in highly regulated, high-stakes environments, I am compelled by Weedmaps' mission to bring transparency and rigorous data standards to the cannabis industry. I see a powerful parallel between my background in managing sensitive, multi-country health data and the challenges of organizing and deriving value from the rapidly evolving cannabis marketplace data.")
    p1.paragraph_format.space_after = Pt(12)
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p2 = document.add_paragraph()
    p2.add_run("In my current role as Lead Biostatistician and Data Strategist for Georgetown University East Africa, I function as the data partner to senior leadership, much like the Staff Data Analyst role calls for. I do not just answer queries; I proactively identify opportunities to optimize multi-million dollar program portfolios. I have successfully architected migrations from legacy reporting tools to modern, automated dashboards (using platforms like Power BI and Looker), acting as a force multiplier that allows non-technical stakeholders to make data-driven decisions independently. This experience aligns directly with your goal of migrating from Tableau to a new BI solution.")
    p2.paragraph_format.space_after = Pt(12)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p3 = document.add_paragraph()
    p3.add_run("Furthermore, I pride myself on being a mentor. I have built and trained data teams from the ground up, establishing code review practices, pair programming sessions, and rigorous data governance standards. I am proficient in the modern data stack—expert in SQL for complex data modeling, comfortable with Python for predictive analytics, and experienced in cloud environments (AWS Certified). I am eager to apply this technical depth to help Weedmaps turn 'messy' data into clear, strategic innovations.")
    p3.paragraph_format.space_after = Pt(12)
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p4 = document.add_paragraph()
    p4.add_run("Thank you for considering my application. I look forward to the possibility of discussing how my experience in driving data adoption and strategic clarity can contribute to Weedmaps' continued growth.")
    p4.paragraph_format.space_after = Pt(24)
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sign off
    p_sign = document.add_paragraph("Sincerely,")
    p_sign.paragraph_format.space_after = Pt(24)

    # Name 
    p_name = document.add_paragraph("Nichodemus Amollo\nnichodemuswerre@gmail.com\n+254 725 737 867")

    document.save('Nichodemus_Amollo_Weedmaps_Cover_Letter.docx')

if __name__ == "__main__":
    create_cl()
