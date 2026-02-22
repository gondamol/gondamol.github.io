from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_cl():
    document = Document()

    # Adjust margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Header Date
    today = datetime.date.today().strftime("%B %d, %Y")
    p_date = document.add_paragraph(today)
    p_date.paragraph_format.space_after = Pt(12)

    # Recipient
    p_recipient = document.add_paragraph()
    p_recipient.add_run("Niki Scotto\nRecruitment Team\nICON plc")
    p_recipient.paragraph_format.space_after = Pt(12)

    # Salutation
    p_salutation = document.add_paragraph("Dear Niki Scotto & Hiring Team,")
    p_salutation.paragraph_format.space_after = Pt(12)

    # Body Paragraphs
    p1 = document.add_paragraph()
    p1.add_run("I am writing to express my strong interest in the Senior Biostatistician role at ICON. As a Senior Statistician and Research Officer at KEMRI and currently the Lead Biostatistician for Georgetown University East Africa, I have spent the last 8+ years immersed in the rigorous world of clinical research and health data. I am eager to bring my experience in Study Design, SAP development, and TLF generation to your diverse team at ICON, contributing to your mission of shaping the future of clinical development.")
    p1.paragraph_format.space_after = Pt(12)
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p2 = document.add_paragraph()
    p2.add_run("The job description emphasizes a need for leadership in study support—from randomization to reporting. In my current role, I lead the statistical architecture for multi-site health interventions, directly overseeing power calculations, stratification schemas, and the development of Statistical Analysis Plans (SAPs). Much like ICON's commitment to quality, I have implemented rigorous data quality control specifications, writing custom R and SAS programs to validate auditing trails ensuring our studies remain audit-ready at all times.")
    p2.paragraph_format.space_after = Pt(12)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p3 = document.add_paragraph()
    p3.add_run("I also deeply value cross-functional collaboration. Whether coordinating with data managers to design Case Report Forms (CRFs) or presenting interim safety results to Principal Investigators, I pride myself on bridging the gap between complex statistical methods and actionable clinical insights. My technical proficiency spans SAS, R, and Python, allowing me to be versatile in any data environment, while my background at KEMRI has grounded me in the strict compliance standards (GCP/ICH) required for pharmaceutical and clinical research.")
    p3.paragraph_format.space_after = Pt(12)
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p4 = document.add_paragraph()
    p4.add_run("I am excited about the opportunity to join ICON's culture of process improvement and look forward to potentially discussing how my statistical leadership can support your clinical projects.")
    p4.paragraph_format.space_after = Pt(24)
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sign off
    p_sign = document.add_paragraph("Sincerely,")
    p_sign.paragraph_format.space_after = Pt(24)

    # Name 
    p_name = document.add_paragraph("Nichodemus Amollo\nnichodemuswerre@gmail.com\n+254 725 737 867")

    document.save('Nichodemus_Amollo_ICON_Cover_Letter.docx')

if __name__ == "__main__":
    create_cl()
