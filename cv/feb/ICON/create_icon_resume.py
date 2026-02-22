from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_resume():
    document = Document()

    # Adjust margins - Standard professional for clinical roles
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # --- CONTENT ---
    
    # Header
    name = "NICHODEMUS AMOLLO"
    title_line = "Senior Biostatistician | Clinical Research Lead"
    contact_line = "Nairobi, Kenya | +254 725 737 867 | nichodemuswerre@gmail.com"
    portfolio_line = "GitHub: github.com/gondamol | LinkedIn: linkedin.com/in/nichodemusamollo"

    summary = (
        "Senior Biostatistician with over 8 years of experience in clinical study design, data analysis, and regulatory "
        "reporting within healthcare and research environments (KEMRI, Georgetown). Proven expertise in developing Statistical "
        "Analysis Plans (SAPs), calculating sample sizes, and conducting rigorous data quality control for multi-site clinical "
        "trials. Proficient in SAS, R, and Python for CDISC-compliant data manipulation and TLF (Tables, Listings, Figures) "
        "generation. Demonstrated leadership in cross-functional collaboration with data management and clinical teams to "
        "ensure data integrity and audit readiness. Committed to driving clinical excellence through statistical rigor and "
        "innovation."
    )

    core_competencies = [
        ("Clinical Study Design", "Sample Size Calculation, Randomization, Protocol Development, Case Report Form (CRF) Design"),
        ("Statistical Analysis", "SAP Development, TLF Generation, Survival Analysis, Regression Modeling, CDISC Standards"),
        ("Data Management", "Data Quality Control (QC), Audit Trails, Data Cleaning & Validation, query resolution"),
        ("Programming", "SAS (Base/Advanced concepts), R (Tidyverse, Survival), Python, SQL"),
        ("Regulatory & Compliance", "GCP Guidelines, ICH Standards, Ethical Research Compliance, Audit Preparation")
    ]

    experience = [
        {
            "role": "Lead Biostatistician",
            "company": "Georgetown University East Africa (gui2de)",
            "location_date": "Nairobi, Kenya | Apr 2025 – Present",
            "details": [
                "**Study Design & Protocol:** Lead the statistical design for multi-site health intervention studies, including power analysis, sample size determination, and stratified randomization schemas.",
                "**Statistical Analysis Plans (SAPs):** Author and review SAPs for longitudinal health studies, defining primary/secondary endpoints and derived datasets in alignment with study protocols.",
                "**Data Quality & QC:** Oversee data validation and quality control procedures, writing SAS/R programs to identify discrepancies and ensure database lock readiness.",
                "**Reporting:** Generate and validate Tables, Listings, and Figures (TLFs) for interim and final clinical study reports (CSRs), presenting findings to Principal Investigators and safety monitoring boards."
            ]
        },
        {
            "role": "Senior Statistician & Research Officer",
            "company": "Kenya Medical Research Institute (KEMRI)",
            "location_date": "Nairobi, Kenya | Apr 2023 – Mar 2025",
            "details": [
                "**Clinical Data Management:** Collaborated with Data Managers to design Case Report Forms (CRFs) and separate CRF completion guidelines, ensuring robust data capture for oncology and infectious disease trials.",
                "**Statistical Programming:** Developed reusable R and SAS macros for standard safety and efficacy analysis, reducing programming time for routine datasets by 40%.",
                "**Audit & Compliance:** Maintained rigorous study documentation and audit trails for all statistical programs, ensuring full compliance with GCP and institutional review board requirements.",
                "**Cross-Functional Collaboration:** Served as the statistical point of contact for clinical operations teams, resolving data queries and providing statistical input for protocol amendments."
            ]
        },
        {
            "role": "Biostatistical Consultant",
            "company": "Various Clinical Research Projects (Global Fund, MMV)",
            "location_date": "Remote / Kenya | 2023 – 2024",
            "details": [
                "**Global Fund PMTCT Study:** Designed the statistical framework and analysis plan for a multi-country HIV effectiveness study, ensuring data harmonization across diverse sites.",
                "**MMV (Medicines for Malaria Venture) Study:** Conducted quantitative analysis on patient uptake data, producing validation reports and statistical summaries for regulatory submissions."
            ]
        },
        {
            "role": "Data Analyst & MEAL Specialist",
            "company": "LERIS Hub",
            "location_date": "Kenya | Sep 2017 – May 2021",
            "details": [
                "**Outcome Analysis:** Conducted survival analysis and logistic regression modeling to evaluate health program outcomes, directly informing adaptive management strategies.",
                "**Data Integrity:** Implemented automated data quality checks using SQL and R, significantly reducing data entry errors in large-scale health surveys."
            ]
        }
    ]

    education = [
        {
            "degree": "MSc – Statistics & Data Science (Biostatistics Focus) (Expected 2026)",
            "school": "Jaramogi Oginga Odinga University of Science & Technology",
            "detail": "Thesis: Financial Determinants & Health Outcomes Modeling" 
        },
        {
            "degree": "BSc (Honours) – Statistics",
            "school": "University of Nairobi",
            "detail": None
        }
    ]

    certs = [
        "**Biomedical Research Ethics (CITI Program)** - GCP & Human Subjects Research",
        "**Monitoring & Evaluation for Global Health** – University of Washington",
        "**Advanced Data Analysis with R** – Johns Hopkins University"
    ]

    tech_skills = [
        "**Statistical Software:** R (Expert), SAS (Proficient), Python (Pandas/Statsmodels), SPSS, Stata.",
        "**Standards & Tools:** CDISC knowledge (SDTM/ADaM concepts), SQL, Git/GitHub, REDCap, ODK.",
        "**Analysis Methods:** Generalized Linear Models, Survival Analysis, Mixed Effects Models, Bayesian Methods."
    ]

    # --- STYLE HELPERS ---
    
    def add_section_header(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        
        # Add border below
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        pbdr = OxmlElement('w:pbdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000') # Black standard border
        pbdr.append(bottom)
        pPr.append(pbdr)

        runner = p.add_run(text.upper())
        runner.bold = True
        runner.font.size = Pt(11)
        runner.font.name = 'Calibri'
        runner.font.color.rgb = RGBColor(0, 0, 0) # Black

    # --- DOCUMENT BUILDING ---

    # HEADER
    p_name = document.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(0)
    r_name = p_name.add_run(name)
    r_name.bold = True
    r_name.font.size = Pt(22)
    r_name.font.name = 'Calibri'
    r_name.font.color.rgb = RGBColor(0, 0, 0) # Black name

    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run(title_line)
    r_title.bold = True
    r_title.font.size = Pt(11)
    r_title.font.name = 'Calibri'

    p_contact = document.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(0)
    r_contact = p_contact.add_run(contact_line)
    r_contact.font.size = Pt(10)
    r_contact.font.name = 'Calibri'

    p_port = document.add_paragraph()
    p_port.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_port.paragraph_format.space_after = Pt(12)
    r_port = p_port.add_run(portfolio_line)
    r_port.font.size = Pt(10)
    r_port.font.name = 'Calibri'
    r_port.font.color.rgb = RGBColor(0, 0, 255) # Blue link

    # PROFESSIONAL SUMMARY
    add_section_header(document, "Professional Summary")
    p_sum = document.add_paragraph(summary)
    p_sum.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_sum.paragraph_format.space_before = Pt(4)
    for run in p_sum.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # TECHNICAL SKILLS
    add_section_header(document, "Technical Skills")
    for skill in tech_skills:
        p_skill = document.add_paragraph(style='List Bullet')
        p_skill.paragraph_format.space_after = Pt(2)
        if "**" in skill:
            parts = skill.split("**")
            for i, part in enumerate(parts):
                r = p_skill.add_run(part)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)
                if i % 2 == 1:
                    r.bold = True
        else:
            r = p_skill.add_run(skill)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

    # CORE COMPETENCIES (Relevant for clinical role)
    add_section_header(document, "Core Competencies")
    for category, content in core_competencies:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_cat = p.add_run(category + ": ")
        r_cat.bold = True
        r_cat.font.name = 'Calibri'
        r_cat.font.size = Pt(11)
        r_con = p.add_run(content)
        r_con.font.name = 'Calibri'
        r_con.font.size = Pt(11)

    # PROFESSIONAL EXPERIENCE
    add_section_header(document, "Professional Experience")
    
    for job in experience:
        # Job Title Line
        p_job = document.add_paragraph()
        p_job.paragraph_format.space_before = Pt(12)
        p_job.paragraph_format.space_after = Pt(0)
        
        r_role = p_job.add_run(job['role'])
        r_role.bold = True
        r_role.font.size = Pt(11)
        r_role.font.name = 'Calibri'
        
        if job['company']:
            p_job.add_run(" | ")
            r_comp = p_job.add_run(job['company'])
            r_comp.bold = True 
            r_comp.font.size = Pt(11)
            r_comp.font.name = 'Calibri'

        # Location | Date Line
        p_loc = document.add_paragraph()
        p_loc.paragraph_format.space_after = Pt(4)
        r_loc = p_loc.add_run(job['location_date'])
        r_loc.italic = True
        r_loc.font.size = Pt(10)
        r_loc.font.name = 'Calibri'

        # Bullets
        for detail in job['details']:
            p_bull = document.add_paragraph(style='List Bullet')
            p_bull.paragraph_format.space_after = Pt(2)
            
            # Formatting logic for **Bold**
            if "**" in detail:
                parts = detail.split("**")
                for i, part in enumerate(parts):
                    r = p_bull.add_run(part)
                    r.font.name = 'Calibri'
                    r.font.size = Pt(11)
                    if i % 2 == 1:
                        r.bold = True
            else:
                r = p_bull.add_run(detail)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)

    # EDUCATION
    add_section_header(document, "Education")
    for edu in education:
        p_edu = document.add_paragraph()
        p_edu.paragraph_format.space_before = Pt(6)
        p_edu.paragraph_format.space_after = Pt(0)
        
        r_deg = p_edu.add_run(edu['degree'])
        r_deg.bold = True
        r_deg.font.name = 'Calibri'
        r_deg.font.size = Pt(11)
        
        p_sch = document.add_paragraph()
        p_sch.paragraph_format.space_after = Pt(0)
        r_sch = p_sch.add_run(edu['school'])
        r_sch.font.name = 'Calibri'
        r_sch.font.size = Pt(11)
        
        if edu['detail']:
            p_det = document.add_paragraph()
            p_det.paragraph_format.space_after = Pt(2)
            r_det = p_det.add_run(edu['detail'])
            r_det.italic = True
            r_det.font.name = 'Calibri'
            r_det.font.size = Pt(10)

    # CERTIFICATIONS
    add_section_header(document, "Certifications")
    for cert in certs:
        p_cert = document.add_paragraph(style='List Bullet')
        p_cert.paragraph_format.space_after = Pt(2)
        if "**" in cert:
            parts = cert.split("**")
            for i, part in enumerate(parts):
                r = p_cert.add_run(part)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)
                if i % 2 == 1:
                    r.bold = True
        else:
            r = p_cert.add_run(cert)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

    document.save('Nichodemus_Amollo_ICON_Resume.docx')

if __name__ == "__main__":
    create_resume()
