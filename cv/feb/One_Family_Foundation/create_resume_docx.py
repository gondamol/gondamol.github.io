from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_resume():
    document = Document()

    # Adjust margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Content
    name = "Nichodemus Amollo"
    title = "Research & Data Lead | Social Innovation Strategist | Health Economics Expert"
    contact = "Nairobi, Kenya | +254 725 737 867 | nichodemuswerre@gmail.com"
    links = "LinkedIn: linkedin.com/in/nichodemusamollo | Portfolio: gondamol.github.io/cv/"

    summary = (
        "Results-oriented Research & Data Lead with over 7 years of experience driving social innovation "
        "through rigorous data science, impact evaluation, and digital product development. Proven track record "
        "in designing and managing complex, multi-country research initiatives—most notably the pioneering "
        "'Health Financial Diaries'—which serve as dynamic engines for understanding poverty dynamics and health financing. "
        "Expert in translating intricate datasets into actionable policy insights (indices, scorecards) for government "
        "and non-profit stakeholders. Passionate about leveraging data to reimagine social systems and empower marginalized communities."
    )

    core_competencies = [
        ("Research Leadership", "Innovation Index Development, RCT Design, Impact Evaluation, Longitudinal Studies (Financial Diaries)"),
        ("Data Management & Analysis", "Advanced Statistical Modeling (R, Python, Stata), Machine Learning, Bayesian Analysis, Big Data Pipelines"),
        ("Social Innovation", "Livelihood Analysis, Financial Behavior Modeling, Health Economics, Policy Strategy"),
        ("Digital Product Oversight", "Dashboard Development (Shiny, Tableau, Power BI), Mobile Data Systems (ODK, SurveyCTO), Automated Reporting"),
        ("Stakeholder Engagement", "Cross-Sector Partnerships (Government, NGOs), Field Team Leadership (50+ staff), Capacity Building & Training")
    ]

    experience = [
        {
            "role": "Lead Research Data Manager & Social Innovation Strategist",
            "company": "Georgetown University Initiative on Innovation, Development and Evaluation (gui2de)",
            "location": "Remote / Kenya & Uganda",
            "dates": "Jan 2020 – Present",
            "details": [
                "**Social Innovation Engine Development:** Architected and led the 'Health Financial Diaries' system—a longitudinal data engine tracking the daily financial lives of 1,000+ low-income households. This 'living product' provides real-time insights into social innovation needs for health financing.",
                "**Research Leadership & Strategy:** Directed end-to-end data operations for 15+ Randomized Controlled Trials (RCTs) and impact evaluations. Defined data taxonomies and structures to measure complex social indicators often invisible in standard surveys.",
                "**Data-Driven Policy Impact:** Transformed raw financial data into 'Financial Health Indices' and policy briefs, directly influencing stakeholder decisions on health insurance products for the informal sector.",
                "**Digital Product Management:** Oversaw the development of automated R-Shiny dashboards that visualize 'Social Innovation' metrics, enabling non-technical stakeholders to monitor project health and impact in real-time.",
                "**Stakeholder & Team Management:** Managed relationships with international PIs, local government bodies, and implementation partners. Led and trained a diverse team of 200+ field researchers, fostering a culture of rigorous data quality and ethical research."
            ]
        },
        {
            "role": "Research Data Analyst | Health Economics & Policy",
            "company": "University of Alaska Anchorage",
            "location": "Remote",
            "dates": "Jan 2019 – Dec 2020",
            "details": [
                "**Policy Research Analysis:** Conducted advanced econometric analysis to identify financial barriers to healthcare access, effectively creating a 'scorecard' for health equity in underserviced regions.",
                "**Data Management:** Managed large-scale datasets to model health expenditure patterns, providing the evidence base for policy recommendations on social protection systems.",
                "**Insight Generation:** Developed interactive visualizations to communicate complex economic findings to diverse audiences, enhancing the visibility of critical social issues."
            ]
        },
        {
            "role": "M&E Specialist & Digital Data Lead",
            "company": "Various Organizations (KEMRI, LERIS Hub, etc.)",
            "location": "Kenya",
            "dates": "Jun 2015 – Dec 2019",
            "details": [
                "**System Design:** Designed and deployed robust Monitoring & Evaluation (M&E) frameworks for health and agriculture programs, focusing on financial sustainability metrics.",
                "**Innovation in Measurement:** Pioneered the use of mobile data collection tools to replace paper-based systems, increasing data accuracy by 40% and reducing report turnaround time.",
                "**Impact Evaluation:** Conducted cost-effectiveness analyses for health interventions, helping organizations optimize their social return on investment (SROI).",
                "**Oncology Data Management:** Built and managed a cancer registry database, contributing to a 'disease burden index' that informed regional health resource allocation."
            ]
        }
    ]

    education = [
        {
            "degree": "MSc in Biostatistics & Epidemiology (In Progress)",
            "school": "Jaramogi Oginga Odinga University of Science and Technology (JOOUST)",
            "details": "Thesis Focus: Financial Determinants & Social Innovation in Chronic Disease Management"
        },
        {
            "degree": "Certificate in Monitoring & Evaluation in Global Health (Expected 2025)",
            "school": "University of Washington",
            "details": ""
        },
        {
            "degree": "BSc in Statistics (2014)",
            "school": "Maseno University, Kenya",
            "details": ""
        }
    ]

    tech_skills = [
        "**Programming & Analysis:** R (Expert - Tidyverse, Shiny), Python (Pandas, Scikit-learn), Stata, SPSS.",
        "**Data Engineering:** SQL, MongoDB, PostgreSQL, API Integration.",
        "**Visualization:** Tableau, Power BI, ggplot2, Interactive Web Dashboards.",
        "**Tools:** Git/GitHub, ODK, SurveyCTO, REDCap, LaTeX/Markdown."
    ]

    achievements = [
        "**Publication:** 'Financial Determinants of Healthcare Seeking Behavior in Rural Kenya' (Journal of Health Economics, In Press).",
        "**Innovation:** Reduced data collection costs by 30% through the implementation of agile, mobile-first data systems.",
        "**Capacity Building:** Trained over 500 researchers across East Africa in advanced data collection and social research methodologies."
    ]

    # --- HELPER FUNCTIONS ---

    def add_section_header(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        runner = p.add_run(text.upper())
        runner.bold = True
        runner.font.size = Pt(11)
        runner.font.color.rgb = RGBColor(0, 51, 102)  # Dark Blue
        # Add a bottom border
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        pbdr = OxmlElement('w:pbdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '003366')
        pbdr.append(bottom)
        pPr.append(pbdr)

    # --- DOCUMENT CONSTRUCTION ---

    # Name
    p_name = document.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(name.upper())
    run_name.bold = True
    run_name.font.size = Pt(20)
    run_name.font.color.rgb = RGBColor(0, 51, 102)

    # Title
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(0)
    run_title = p_title.add_run(title)
    run_title.font.size = Pt(10)
    run_title.bold = True

    # Contact
    p_contact = document.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(0)
    run_contact = p_contact.add_run(contact)
    run_contact.font.size = Pt(9)

    # Links
    p_links = document.add_paragraph()
    p_links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_links.paragraph_format.space_after = Pt(12)
    run_links = p_links.add_run(links)
    run_links.font.size = Pt(9)
    run_links.font.color.rgb = RGBColor(0, 102, 204)

    # Professional Summary
    add_section_header(document, "Professional Summary")
    p_summary = document.add_paragraph(summary)
    p_summary.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Core Competencies
    add_section_header(document, "Core Competencies")
    for category, skills in core_competencies:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        run_cat = p.add_run(f"{category}: ")
        run_cat.bold = True
        p.add_run(skills)

    # Experience
    add_section_header(document, "Professional Experience")
    for job in experience:
        # Role & Company
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(8)
        
        run_role = p.add_run(job['role'])
        run_role.bold = True
        run_role.font.size = Pt(11)
        
        # Use tabs for layout? Or just a formatted line. Let's do formatted line.
        p.add_run(f" | {job['company']}")
        
        # Date & Location line
        p_sub = document.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(4)
        run_loc = p_sub.add_run(f"{job['location']} | {job['dates']}")
        run_loc.italic = True
        run_loc.font.size = Pt(9)

        # Details
        for detail in job['details']:
            p_det = document.add_paragraph(style='List Bullet')
            # Rudimentary bold parsing
            if "**" in detail:
               parts = detail.split("**")
               # Assumption: starts with empty string if ** is at start.
               # e.g. "**Title:** desc" -> ["", "Title:", " desc"]
               for i, part in enumerate(parts):
                   r = p_det.add_run(part)
                   if i % 2 == 1: # Odd indices are between ** **
                       r.bold = True
            else:
                p_det.add_run(detail)

    # Education
    add_section_header(document, "Education")
    for edu in education:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(4)
        r_deg = p.add_run(edu['degree'])
        r_deg.bold = True
        
        p2 = document.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.add_run(edu['school'])

        if edu['details']:
             p3 = document.add_paragraph()
             p3.paragraph_format.space_after = Pt(2)
             r_det = p3.add_run(edu['details'])
             r_det.italic = True
             r_det.font.size = Pt(9)

    # Skills
    add_section_header(document, "Technical Skills")
    for skill in tech_skills:
        p = document.add_paragraph(style='List Bullet')
        if "**" in skill:
           parts = skill.split("**")
           for i, part in enumerate(parts):
               r = p.add_run(part)
               if i % 2 == 1:
                   r.bold = True
        else:
            p.add_run(skill)

    # Achievements
    add_section_header(document, "Selected Achievements")
    for ach in achievements:
        p = document.add_paragraph(style='List Bullet')
        if "**" in ach:
            parts = ach.split("**")
            for i, part in enumerate(parts):
                r = p.add_run(part)
                if i % 2 == 1:
                    r.bold = True
        else:
            p.add_run(ach)

    document.save('Nichodemus_Amollo_Resume.docx')

if __name__ == "__main__":
    create_resume()
