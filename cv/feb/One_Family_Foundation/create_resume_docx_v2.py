from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_resume():
    document = Document()

    # Adjust margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Content
    name = "Nichodemus Amollo"
    title = "Research & Data Lead | Social Innovation Strategist"
    contact = "Nairobi, Kenya | +254 725 737 867 | nichodemuswerre@gmail.com"
    links = "LinkedIn: linkedin.com/in/nichodemusamollo | Portfolio: gondamol.github.io/cv/"

    summary = (
        "I am a Research & Data Lead with over 7 years of experience in social innovation, impact evaluation, "
        "and health economics. My career is defined by a passion for using data to tell human stories and drive "
        "systemic change. From leading the technical design of the 'Health Financial Diaries' to managing multi-country "
        "RCTs, I specialize in building 'living data engines' that inform policy and empower communities. "
        "I bring a unique blend of high-level statistical expertise and on-the-ground field leadership, ensuring that "
        "every data point represents a real life and every insight leads to real impact."
    )

    experience = [
        {
            "role": "Lead Research Data Manager & Strategist",
            "company": "Georgetown University Initiative on Innovation, Development and Evaluation (gui2de)",
            "location": "Remote / Kenya",
            "dates": "Feb 2025 – Present",
            "details": [
                "Leading the data architecture for the 'Health Financial Diaries' project, tracking financial resilience across 1,000+ households.",
                "Designing and deploying automated R-Shiny dashboards to visualize social metrics in real time for policymakers.",
                "Managing relationships with international Principal Investigators and local government stakeholders to robustly effect policy changes.",
                "Directing a diverse team of field researchers, ensuring high-quality data collection through ethical and rigorous protocols."
            ]
        },
        {
            "role": "CONSULTANCY EXPERIENCE",
            "company": "", 
            "location": "",
            "dates": "2023 – 2024", 
            "is_section_header": True,
            "sub_roles": [
                {
                    "role": "Technical Lead",
                    "project": "Global Fund PMTCT Effectiveness Study",
                    "date": "2024",
                    "details": [
                         "Designed the full study protocol and statistical framework.",
                         "Led the data analysis and interpretation of results.",
                         "Supported ethics submission and donor reporting."
                    ]
                },
                {
                    "role": "Lead Consultant",
                    "project": "RAS Uptake Study, Dalberg and MMV",
                    "date": "2024",
                    "details": [
                        "Conducted both quantitative and qualitative analysis.",
                        "Delivered the final analytical report used for strategic decision making."
                    ]
                },
                {
                    "role": "MEL Consultant",
                    "project": "Community Health Dashboards, NEPHAK",
                    "date": "2023",
                    "details": [
                        "Developed automated analytical dashboards for health tracking.",
                        "Supported indicator development and monitoring systems."
                    ]
                }
            ]
        },
        {
            "role": "Graduate Research Fellow",
            "company": "Jaramogi Oginga Odinga University of Science and Technology (JOOUST)",
            "location": "Kenya",
            "dates": "Jan 2021 – Dec 2022",
            "details": [
                "Conducted advanced research on financial determinants of chronic disease management.",
                "Mentored junior students in statistical methods and data collection ethics.",
                "Co-authored papers on health financing and social protection systems."
            ]
        },
        {
            "role": "Research Data Analyst",
            "company": "University of Alaska Anchorage",
            "location": "Remote",
            "dates": "Jan 2019 – Dec 2020",
            "details": [
                "Conducted econometric analysis to identify financial barriers to healthcare access.",
                "Developed interactive visualizations to communicate complex economic findings."
            ]
        },
        {
            "role": "Research Data Specialist",
            "company": "Kenya Medical Research Institute (KEMRI)",
            "location": "Kenya",
            "dates": "Jan 2017 – Dec 2018",
            "details": [
                "Built and managed a cancer registry database for epidemiological research.",
                "Contributed to data quality assessments and disease burden reporting."
            ]
        },
         {
            "role": "Monitoring & Evaluation Officer",
            "company": "LERIS Hub",
            "location": "Kenya",
            "dates": "Jun 2015 – Dec 2016",
            "details": [
                "Designed M&E frameworks for health and agriculture programs.",
                "Pioneered the transition from paper-based to mobile data collection systems."
            ]
        }
    ]

    education = [
        {
            "degree": "MSc in Biostatistics & Epidemiology (In Progress)",
            "school": "Jaramogi Oginga Odinga University (JOOUST)",
            "details": "Thesis Focus: Financial Determinants in Chronic Disease Management"
        },
        {
            "degree": "BSc in Statistics",
            "school": "Maseno University, Kenya",
            "details": "Graduated 2014"
        }
    ]
    
    certifications = [
        "Certificate in Monitoring & Evaluation in Global Health (University of Washington)",
        "Social Innovation Management (Amani Institute)",
        "Google Data Analytics Professional Certificate",
        "Human-Centered Design for Social Innovation (Acumen)",
        "Advanced Data Analysis with R (Johns Hopkins University)"
    ]

    tech_skills = [
        "**Analysis & Coding:** R (Expert), Python, Stata, SPSS",
        "**Data Engineering:** SQL, MongoDB, PostgreSQL, ODK, SurveyCTO",
        "**Visualization:** Tableau, Power BI, R-Shiny, ggplot2",
        "**Tools:** Git, GitHub, LaTeX, RedCap"
    ]

    # --- HELPER FUNCTIONS ---

    def add_section_header(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        runner = p.add_run(text.upper())
        runner.bold = True
        runner.font.size = Pt(11)
        runner.font.color.rgb = RGBColor(0, 51, 102)  # Dark Blue
        # Add a bottom border
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        pbdr = OxmlElement('w:pbdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '003366')
        pbdr.append(bottom)
        pPr.append(pbdr)

    # --- DOCUMENT CONSTRUCTION ---

    # Name
    p_name = document.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(name.upper())
    run_name.bold = True
    run_name.font.size = Pt(20)
    run_name.font.color.rgb = RGBColor(0, 51, 102)

    # Title
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(0)
    run_title = p_title.add_run(title)
    run_title.font.size = Pt(10)
    run_title.bold = True

    # Contact
    p_contact = document.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(0)
    run_contact = p_contact.add_run(contact)
    run_contact.font.size = Pt(9)

    # Links
    p_links = document.add_paragraph()
    p_links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_links.paragraph_format.space_after = Pt(12)
    run_links = p_links.add_run(links)
    run_links.font.size = Pt(9)
    run_links.font.color.rgb = RGBColor(0, 102, 204)

    # Professional Summary
    add_section_header(document, "Professional Summary")
    p_summary = document.add_paragraph(summary)
    p_summary.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Experience
    add_section_header(document, "Professional Experience")
    for job in experience:
        if job.get("is_section_header"):
            # Sub-header for Consultancies
            p_head = document.add_paragraph()
            p_head.paragraph_format.space_before = Pt(10)
            p_head.paragraph_format.space_after = Pt(4)
            r_head = p_head.add_run(job["role"])
            r_head.bold = True
            r_head.font.size = Pt(11)
            r_head.font.color.rgb = RGBColor(0, 51, 102)
            
            p_dates = document.add_paragraph()
            p_dates.paragraph_format.space_after = Pt(6)
            r_dates = p_dates.add_run(job["dates"])
            r_dates.italic = True
            
            for sub in job["sub_roles"]:
                p_sub = document.add_paragraph()
                p_sub.paragraph_format.space_before = Pt(4)
                p_sub.paragraph_format.left_indent = Inches(0.2)
                
                # Role and Project line
                r_role = p_sub.add_run(sub["role"])
                r_role.bold = True
                p_sub.add_run(f" | {sub['project']} ({sub['date']})")
                
                # Details
                for det in sub["details"]:
                    p_det = document.add_paragraph(style='List Bullet')
                    p_det.paragraph_format.left_indent = Inches(0.4)
                    p_det.add_run(det)

        else:
            # Standard Role
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(8)
            
            run_role = p.add_run(job['role'])
            run_role.bold = True
            run_role.font.size = Pt(11)
            
            p.add_run(f" | {job['company']}")
            
            # Date & Location line
            p_sub = document.add_paragraph()
            p_sub.paragraph_format.space_after = Pt(4)
            run_loc = p_sub.add_run(f"{job['location']} | {job['dates']}")
            run_loc.italic = True
            run_loc.font.size = Pt(9)

            # Details
            for detail in job['details']:
                p_det = document.add_paragraph(style='List Bullet')
                p_det.add_run(detail)

    # Education
    add_section_header(document, "Education")
    for edu in education:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(4)
        r_deg = p.add_run(edu['degree'])
        r_deg.bold = True
        
        p2 = document.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.add_run(edu['school'])

        if edu['details']:
             p3 = document.add_paragraph()
             p3.paragraph_format.space_after = Pt(2)
             r_det = p3.add_run(edu['details'])
             r_det.italic = True
             r_det.font.size = Pt(9)
             
    # Certifications
    add_section_header(document, "Certifications")
    for cert in certifications:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(cert)

    # Skills
    add_section_header(document, "Technical Skills")
    for skill in tech_skills:
        p = document.add_paragraph(style='List Bullet')
        if "**" in skill:
           parts = skill.split("**")
           for i, part in enumerate(parts):
               r = p.add_run(part)
               if i % 2 == 1:
                   r.bold = True
        else:
            p.add_run(skill)

    document.save('Nichodemus_Amollo_Resume.docx')

if __name__ == "__main__":
    create_resume()
