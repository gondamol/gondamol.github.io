from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_resume():
    document = Document()

    # Adjust margins - narrow for more space
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Content Data
    name = "NICHODEMUS AMOLLO"
    title_line = "Research & Data Lead | Health Economics Expert"
    contact_line = "Nairobi, Kenya | +254 725 737 867 | nichodemuswerre@gmail.com"
    portfolio_line = "Portfolio: gondamol.github.io/cv/"

    summary = (
        "Results-oriented Research & Data Lead with over 7 years of experience driving social innovation "
        "through rigorous data science, impact evaluation, and digital product development. Proven track record "
        "in designing and managing complex, multi-country research initiatives, most notably the pioneering "
        "'Health Financial Diaries'—which serve as dynamic engines for understanding poverty dynamics and health financing. "
        "Expert in translating intricate datasets into actionable policy insights (indices, scorecards) for government "
        "and non-profit stakeholders. Passionate about leveraging data to reimagine social systems and empower marginalized communities."
    )

    core_competencies = [
        ("Research Leadership", "Innovation Index Development, RCT Design, Impact Evaluation, Longitudinal Studies (Financial Diaries)"),
        ("Data Management & Analysis", "Advanced Statistical Modeling (R, Python, Stata), Machine Learning, Bayesian Analysis, Big Data Pipelines"),
        ("Social Innovation", "Livelihood Analysis, Financial Behavior Modeling, Health Economics, Policy Strategy"),
        ("Digital Product Oversight", "Dashboard Development (Shiny, Tableau, Power BI), Mobile Data Systems (ODK, SurveyCTO), Automated Reporting"),
        ("Stakeholder Engagement", "Cross-Sector Partnerships (Government, NGOs), Field Team Leadership (50+ staff), Capacity Building & Training")
    ]

    experience = [
        {
            "role": "Lead Research Data Manager",
            "company": "Georgetown University Initiative on Innovation, Development and Evaluation (gui2de)",
            "location_date": "Remote / Kenya & Uganda | Feb 2025 – Present",
            "details": [
                "**Social Innovation Engine Development:** Architected and led the 'Health Financial Diaries' system, a longitudinal data engine tracking the daily financial lives of 1,000+ low-income households. This 'living product' provides real-time insights into social innovation needs for health financing.",
                "**Research Leadership & Strategy:** Directed end-to-end data operations for 8+ Randomized Controlled Trials (RCTs) and impact evaluations. Defined data taxonomies and structures to measure complex social indicators often invisible in standard surveys.",
                "**Data-Driven Policy Impact:** Transformed raw financial data into 'Financial Health Indices' and policy briefs, directly influencing stakeholder decisions on health insurance products for the informal sector.",
                "**Digital Product Management:** Oversaw the development of automated R-Shiny dashboards that visualize 'Social Innovation' metrics, enabling non-technical stakeholders to monitor project health and impact in real-time.",
                "**Stakeholder & Team Management:** Managed relationships with international PIs, local government bodies, and implementation partners. Led and trained a diverse team of 200+ field researchers, fostering a culture of rigorous data quality and ethical research."
            ]
        },
        {
            "role": "Research Data Analyst | Health Economics & Policy",
            "company": "University of Alaska Anchorage",
            "location_date": "Remote | Jan 2022 – Dec 2024", 
            "details": [
                "**Policy Research Analysis:** Conducted advanced econometric analysis to identify financial barriers to healthcare access, effectively creating a 'scorecard' for health equity in underserviced regions.",
                "**Data Management:** Managed large-scale datasets to model health expenditure patterns, providing the evidence base for policy recommendations on social protection systems.",
                "**Insight Generation:** Developed interactive visualizations to communicate complex economic findings to diverse audiences, enhancing the visibility of critical social issues."
            ]
        },
        {
            "role": "M&E Specialist & Digital Data Lead",
            "company": "(KEMRI, LERIS Hub, etc.)",
            "location_date": "Kenya | Jun 2015 – Dec 2021",
            "details": [
                "**System Design:** Designed and deployed robust Monitoring & Evaluation (M&E) frameworks for health and agriculture programs, focusing on financial sustainability metrics.",
                "**Innovation in Measurement:** Pioneered the use of mobile data collection tools to replace paper-based systems, increasing data accuracy by 40% and reducing report turnaround time.",
                "**Impact Evaluation:** Conducted cost-effectiveness analyses for health interventions, helping organizations optimize their social return on investment (SROI).",
                "**Oncology Data Management:** Built and managed a cancer registry database, contributing to a 'disease burden index' that informed regional health resource allocation."
            ]
        }
    ]

    education = [
        {
            "degree": "MSc in Biostatistics & Epidemiology (In Progress)",
            "school": "Jaramogi Oginga Odinga University of Science and Technology (JOOUST)",
            "detail": "Thesis Focus: Health Financing in Chronic Disease Management"
        },
        {
            "degree": "Certificate in Monitoring & Evaluation in Global Health (Expected 2025)",
            "school": "University of Washington",
            "detail": None
        },
        {
            "degree": "BSc in Statistics (2016)",
            "school": "University of Nairobi, Kenya",
            "detail": None
        }
    ]

    tech_skills = [
        "**Programming & Analysis:** R (Expert - Tidyverse, Shiny), Python (Pandas, Scikit-learn), Stata, SPSS.",
        "**Data Engineering:** SQL, MongoDB, PostgreSQL, API Integration.",
        "**Visualization:** Tableau, Power BI, ggplot2, Interactive Web Dashboards.",
        "**Tools:** Git/GitHub, ODK, SurveyCTO, REDCap, LaTeX/Markdown."
    ]

    achievements = [
        "**Publication:** 'Financial Determinants of Effective Hypertension and Diabetes Care in Rural Primary Health Facilities in Kisumu, Kenya: A Mixed-Methods Study (BMC Journal of Health Economics, In Press).",
        "**Innovation:** Reduced data collection costs by 30% through the implementation of agile, mobile-first data systems.",
        "**Capacity Building:** Trained over 500 researchers across East Africa in advanced data collection and social research methodologies."
    ]

    # --- STYLE HELPERS ---
    
    def add_section_header(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        
        # Add border below
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        pbdr = OxmlElement('w:pbdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000') # Black standard border
        pbdr.append(bottom)
        pPr.append(pbdr)

        runner = p.add_run(text.upper())
        runner.bold = True
        runner.font.size = Pt(11)
        runner.font.name = 'Calibri'
        runner.font.color.rgb = RGBColor(0, 0, 0) # Black

    # --- DOCUMENT BUILDING ---

    # HEADER
    p_name = document.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(0)
    r_name = p_name.add_run(name)
    r_name.bold = True
    r_name.font.size = Pt(22)
    r_name.font.name = 'Calibri'
    r_name.font.color.rgb = RGBColor(0, 0, 0) # Black name

    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run(title_line)
    r_title.bold = True
    r_title.font.size = Pt(11)
    r_title.font.name = 'Calibri'

    p_contact = document.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(0)
    r_contact = p_contact.add_run(contact_line)
    r_contact.font.size = Pt(10)
    r_contact.font.name = 'Calibri'

    p_port = document.add_paragraph()
    p_port.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_port.paragraph_format.space_after = Pt(12)
    r_port = p_port.add_run(portfolio_line)
    r_port.font.size = Pt(10)
    r_port.font.name = 'Calibri'
    r_port.font.color.rgb = RGBColor(0, 0, 255) # Blue link

    # PROFESSIONAL SUMMARY
    add_section_header(document, "Professional Summary")
    p_sum = document.add_paragraph(summary)
    p_sum.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_sum.paragraph_format.space_before = Pt(4)
    for run in p_sum.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # CORE COMPETENCIES
    add_section_header(document, "Core Competencies")
    for category, content in core_competencies:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.0) # Flush left
        
        r_cat = p.add_run(category + ": ")
        r_cat.bold = True
        r_cat.font.name = 'Calibri'
        r_cat.font.size = Pt(11)
        
        r_con = p.add_run(content)
        r_con.font.name = 'Calibri'
        r_con.font.size = Pt(11)

    # PROFESSIONAL EXPERIENCE
    add_section_header(document, "Professional Experience")
    
    for job in experience:
        # Job Title Line
        p_job = document.add_paragraph()
        p_job.paragraph_format.space_before = Pt(12)
        p_job.paragraph_format.space_after = Pt(0)
        
        r_role = p_job.add_run(job['role'])
        r_role.bold = True
        r_role.font.size = Pt(11)
        r_role.font.name = 'Calibri'
        
        if job['company']:
            p_job.add_run(" | ")
            r_comp = p_job.add_run(job['company'])
            r_comp.bold = True # Company also bold or normal? Usually beneficial to keep bold or distinct
            r_comp.font.size = Pt(11)
            r_comp.font.name = 'Calibri'

        # Location | Date Line
        p_loc = document.add_paragraph()
        p_loc.paragraph_format.space_after = Pt(4)
        r_loc = p_loc.add_run(job['location_date'])
        r_loc.italic = True
        r_loc.font.size = Pt(10) # Slightly smaller
        r_loc.font.name = 'Calibri'

        # Bullets
        for detail in job['details']:
            p_bull = document.add_paragraph(style='List Bullet')
            p_bull.paragraph_format.space_after = Pt(2)
            
            # Formatting logic for **Bold**
            if "**" in detail:
                parts = detail.split("**")
                for i, part in enumerate(parts):
                    r = p_bull.add_run(part)
                    r.font.name = 'Calibri'
                    r.font.size = Pt(11)
                    if i % 2 == 1:
                        r.bold = True
            else:
                r = p_bull.add_run(detail)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)

    # EDUCATION
    add_section_header(document, "Education")
    for edu in education:
        p_edu = document.add_paragraph()
        p_edu.paragraph_format.space_before = Pt(6)
        p_edu.paragraph_format.space_after = Pt(0)
        
        r_deg = p_edu.add_run(edu['degree'])
        r_deg.bold = True
        r_deg.font.name = 'Calibri'
        r_deg.font.size = Pt(11)
        
        p_sch = document.add_paragraph()
        p_sch.paragraph_format.space_after = Pt(0)
        r_sch = p_sch.add_run(edu['school'])
        r_sch.font.name = 'Calibri'
        r_sch.font.size = Pt(11)
        
        if edu['detail']:
            p_det = document.add_paragraph()
            p_det.paragraph_format.space_after = Pt(2)
            r_det = p_det.add_run(edu['detail'])
            r_det.italic = True
            r_det.font.name = 'Calibri'
            r_det.font.size = Pt(10)

    # TECHNICAL SKILLS
    add_section_header(document, "Technical Skills")
    for skill in tech_skills:
        p_skill = document.add_paragraph(style='List Bullet')
        p_skill.paragraph_format.space_after = Pt(2)
        if "**" in skill:
            parts = skill.split("**")
            for i, part in enumerate(parts):
                r = p_skill.add_run(part)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)
                if i % 2 == 1:
                    r.bold = True
        else:
            r = p_skill.add_run(skill)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

    # ACHIEVEMENTS
    add_section_header(document, "Selected Achievements")
    for ach in achievements:
        p_ach = document.add_paragraph(style='List Bullet')
        p_ach.paragraph_format.space_after = Pt(2)
        if "**" in ach:
            parts = ach.split("**")
            for i, part in enumerate(parts):
                r = p_ach.add_run(part)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)
                if i % 2 == 1:
                    r.bold = True
        else:
            r = p_ach.add_run(ach)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

    document.save('Nichodemus_Amollo_Resume_Refined.docx')

if __name__ == "__main__":
    create_resume()
