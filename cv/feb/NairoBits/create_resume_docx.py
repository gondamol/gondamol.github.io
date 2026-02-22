from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_resume_docx():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run('NICHODEMUS WERRE AMOLLO')
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    
    subheader = doc.add_paragraph()
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role_run = subheader.add_run('Monitoring, Evaluation & Learning (MEL) Specialist | Data & Impact Analyst')
    role_run.bold = True
    role_run.font.size = Pt(12)
    role_run.font.color.rgb = RGBColor(14, 165, 233) # Sky 500
    subheader.paragraph_format.space_after = Pt(6)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = 'Nairobi, Kenya | +254 725 737 867 | nichodemuswerre@gmail.com'
    contact.add_run(contact_text)
    
    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link_run = links.add_run('LinkedIn: linkedin.com/in/nichodemusamollo | Portfolio: gondamol.github.io/cv/')
    link_run.font.size = Pt(10)
    links.paragraph_format.space_after = Pt(18)

    def add_section_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(15, 23, 42)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        # Add bottom border simulation (underline)
        p_border = doc.add_paragraph()
        run_border = p_border.add_run('━' * 65) # Simple border
        run_border.font.size = Pt(6)
        run_border.font.color.rgb = RGBColor(200, 200, 200)
        p_border.paragraph_format.space_after = Pt(12)
        p_border.paragraph_format.line_spacing = Pt(2)
        # Using a solid line graphic or proper borders in pure python-docx is tricky, 
        # so relying on clear typography instead.

    def add_job_header(role, company, date):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        
        # Table for layout
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Column widths
        table.columns[0].width = Inches(5.0)
        table.columns[1].width = Inches(1.5)
        
        row = table.rows[0]
        c1 = row.cells[0].paragraphs[0]
        c2 = row.cells[1].paragraphs[0]
        
        r1 = c1.add_run(role)
        r1.bold = True
        r1.font.size = Pt(12)
        
        r2 = c1.add_run('\n' + company)
        r2.italic = True
        r2.font.color.rgb = RGBColor(80, 80, 80)
        
        r3 = c2.add_run(date)
        r3.italic = True
        c2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Professional Summary
    add_section_header('Professional Summary')
    doc.add_paragraph(
        "Results-oriented MEL Specialist with over 8 years of experience in designing and executing robust monitoring, "
        "evaluation, and learning frameworks for impactful development programs. Expert in harnessing advanced data analytics "
        "(SQL, Python, Stata) to drive strategic decision-making, optimize program performance, and demonstrate tangible impact. "
        "Adept at managing complex data collection systems, building staff capacity, and producing high-quality donor reports. "
        "Committed to leveraging evidence-based insights to empower underserved communities and advance digital inclusion."
    )

    # Areas of Expertise
    add_section_header('Areas of Expertise')
    expertise_p = doc.add_paragraph()
    expertise_p.paragraph_format.line_spacing = 1.3
    expertise_text = (
        "• M&E Framework Design: Results Frameworks, Logframes, Theory of Change, KPI Development\n"
        "• Data Management & Analysis: ODK, SurveyCTO, KoboToolbox, SQL, Python, Stata, R\n"
        "• Impact Reporting: Detailed Narrative Reports, Donor Reporting, Data Visualization (Tableau, PowerBI)\n"
        "• Strategic Planning: Baseline/Endline Surveys, Risk Management, Impact Evaluation Design\n"
        "• Capacity Building: Training Facilitation, Data Quality Assurance, Mentorship\n"
        "• Sector Knowledge: Youth Empowerment, Digital Skills, Health Systems, Financial Inclusion"
    )
    expertise_p.add_run(expertise_text)

    # Professional Experience
    add_section_header('Professional Experience')

    # Job 1
    add_job_header('Lead Research Data Scientist & MEL Specialist', 'Georgetown University (gui2de)', '2023 – Present')
    p1 = doc.add_paragraph()
    p1.style = 'List Bullet'
    p1.add_run("M&E Systems Design: Architected and implemented a high-frequency M&E system (\"Financial Diaries\") tracking economic resilience indicators for 1,000+ households.")
    p2 = doc.add_paragraph()
    p2.style = 'List Bullet'
    p2.add_run("Data Quality Assurance: Developed and operationalized comprehensive data quality protocols, reducing discrepancies by 30%.")
    p3 = doc.add_paragraph()
    p3.style = 'List Bullet'
    p3.add_run("Impact Reporting: Synthesized complex data into actionable insights, directly influencing intervention strategies.")
    p4 = doc.add_paragraph()
    p4.style = 'List Bullet'
    p4.add_run("Strategic Support: Provided critical data inputs for quarterly donor reports, demonstrating program impact.")
    p5 = doc.add_paragraph()
    p5.style = 'List Bullet'
    p5.add_run("Capacity Building: Mentored a team of 5 junior analysts in data management best practices.")

    # Job 2
    add_job_header('Senior Research Data Scientist (M&E Lead)', 'KEMRI (Medical Research Institute)', '2020 – 2023')
    j2_items = [
        "Innovative Monitoring Tools: Spearheaded the design of an \"Early Warning System\" reducing response times from 2 weeks to 3 days.",
        "Cross-Functional Coordination: Coordinated with field teams to harmonize data collection tools across multiple project sites.",
        "Reporting & Communication: Produced comprehensive impact reports for international stakeholders detailing achievements and lessons learned.",
        "Data Utilization: Conducted regular data audits and feedback sessions with field staff to improve data collection standards."
    ]
    for item in j2_items:
        p = doc.add_paragraph(item, style='List Bullet')

    # Job 3
    add_job_header('Evaluation Lead', 'LERIS Hub', '2017 – 2020')
    j3_items = [
        "Program Evaluation: Led M&E of youth empowerment programs, overseeing baseline, mid-term, and end-line surveys.",
        "Inclusive Data Collection: Implemented \"Offline-First\" mobile systems to reach marginalized rural populations, increasing coverage by 25%.",
        "Donor Reporting: Authored detailed narrative reports highlighting program success in increasing digital literacy.",
        "Capacity Building: Facilitated M&E workshops for program staff to enhance data utilization."
    ]
    for item in j3_items:
        p = doc.add_paragraph(item, style='List Bullet')

    # Education
    add_section_header('Education')
    edu_p = doc.add_paragraph()
    edu_p.add_run('MSc in Statistics and Data Science').bold = True
    edu_p.add_run(' | Jaramogi Oginga Odinga University (JOOUST) (Expected 2026)')
    
    edu_p2 = doc.add_paragraph()
    edu_p2.add_run('Graduate Certificate, Social Impact Analytics').bold = True
    edu_p2.add_run(' | University of San Francisco (2025)')

    edu_p3 = doc.add_paragraph()
    edu_p3.add_run('BSc in Statistics').bold = True
    edu_p3.add_run(' | Maseno University (2016)')

    # Certifications
    add_section_header('Certifications')
    cert_p = doc.add_paragraph()
    cert_p.add_run('Project Management (Google) | Research Methodology (Johns Hopkins) | Advanced Data Analysis (DataCamp)')

    doc.save('c:/Users/HFD 2/pc/JBS/2026/feb/NairoBits/Nichodemus_Amollo_Resume.docx')
    print("Resume generated successfully.")

if __name__ == "__main__":
    create_resume_docx()
