from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cover_letter():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header / Contact Info
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run('NICHODEMUS WERRE AMOLLO')
    name_run.bold = True
    name_run.font.size = Pt(14)
    name_run.font.color.rgb = None # Default black
    
    contact = doc.add_paragraph('Nairobi, Kenya | +254 725 737 867 | nichodemuswerre@gmail.com')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(24)
    
    # Date
    doc.add_paragraph('14th February 2026')
    
    # Recipient
    recipient = doc.add_paragraph()
    recipient.add_run('Hiring Manager\n').bold = True
    recipient.add_run('NairoBits Trust\nNairobi, Kenya')
    recipient.paragraph_format.space_after = Pt(12)

    # Subject Line
    subject = doc.add_paragraph()
    subject_run = subject.add_run('RE: APPLICATION FOR MONITORING, EVALUATION AND LEARNING (MEL) ASSOCIATE')
    subject_run.bold = True
    subject_run.underline = True
    subject.paragraph_format.space_after = Pt(12)

    # Salutation
    doc.add_paragraph('Dear Hiring Manager,')

    # Body Paragraphs
    body_1 = doc.add_paragraph(
        "I am writing to express my strong interest in the Monitoring, Evaluation and Learning (MEL) Associate position at NairoBits Trust. "
        "As a Data Scientist and MEL Specialist with over 8 years of experience in the development sector, "
        "my career has been defined by a single mission: using data to ensure that interventions truly transform lives. "
        "I am living proof of the power of targeted support—having transitioned from a beneficiary of aid programs to designing the very systems that evaluate them. "
        "I am eager to bring this personal drive and professional expertise to NairoBits, ensuring that your digital skills training programs "
        "continue to empower youth with the same transformative opportunities that shaped my own path."
    )
    body_1.paragraph_format.space_after = Pt(12)

    body_2 = doc.add_paragraph(
        "In my current role as Lead Research Data Scientist at Georgetown University (gui2de), I designed and executed a high-frequency \"Financial Diaries\" "
        "M&E system tracking the economic resilience of over 1,000 households. This was not merely a data collection exercise; "
        "it was a strategic tool that allowed us to visualize \"invisible\" vulnerabilities and pivot our programming in real-time. "
        "This experience directly aligns with your need for a candidate who can design effective monitoring frameworks and tools. "
        "I am expert in translating complex logframes into operational workflows using ODK and SurveyCTO, ensuring that data from the field—whether "
        "from training centers or community outreaches—is accurate, authorized, and actionable."
    )
    body_2.paragraph_format.space_after = Pt(12)
    
    body_3 = doc.add_paragraph(
        "Beyond the technical design of M&E systems, I am deeply committed to capacity building and organizational learning. "
        "At KEMRI, I led cross-functional teams to harmonize data collection efforts across multiple sites, turning fragmented data "
        "into a cohesive \"Early Warning System\" that reduced response times from weeks to days. I take pride in mentoring program staff, "
        "helping them transition from passive data collectors to active users of information for strategic planning. "
        "My advanced skills in Python, SQL, and Stata allow me to produce the kind of rigorous, high-quality impact reports "
        "that not only satisfy donor requirements but also tell a compelling story of change."
    )
    body_3.paragraph_format.space_after = Pt(12)

    body_4 = doc.add_paragraph(
        "NairoBits’ mission to place youth at the center of societal solutions resonates with me. "
        "Having led evaluation efforts for youth empowerment programs at LERIS Hub, where I implemented \"Offline-First\" tools to reach marginalized rural populations, "
        "I understand the unique challenges and opportunities in this space. I am ready to apply my skills in risk management, impact reporting, "
        "and strategic M&E to support NairoBits in equipping the next generation with future-ready skills."
    )
    body_4.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        "Thank you for considering my application. I would welcome the opportunity to discuss how my blend of technical precision and practical field experience "
        "can advance NairoBits Trust’s impactful work."
    )
    
    # Sign-off
    doc.add_paragraph('Sincerely,')
    
    signature = doc.add_paragraph()
    signature.paragraph_format.space_before = Pt(24)
    sign_run = signature.add_run('Nichodemus Werre Amollo')
    sign_run.bold = True

    # Save
    doc.save('c:/Users/HFD 2/pc/JBS/2026/feb/NairoBits/Nichodemus_Amollo_Cover_Letter.docx')
    print("Cover letter generated successfully.")

if __name__ == "__main__":
    create_cover_letter()
